import requests
import re
import nltk
from nltk.corpus import stopwords
import os

# Descargar recursos necesarios para NLTK
nltk.download('punkt')
nltk.download('stopwords')

os.system('\n\npause')
os.system('cls')

print('\n■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■ INICIO del Programa •Documento de WORD - VALIDACION• ■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■')

Programa=[
    '\n○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○ M E N U ○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○\n',
    '■ Realizar las siguientes acciones sobre un documento de WORD ■',
    '► Ingresar un Archivo *.DOCX almacenado en la computadora ("Validar su Ubicación y su formato").',
    '•01• Enviar el texto del archivo *.DOCX a un archivo de texto.txt almacenado en la computadora.',
    '•02• Contar  el numero de palabras.',
    '•03• Contar  el numero de lineas.',
    '•04• Mostrar palabras de 3 o 4 caracteres.',
    '•05• Cuenta el numero de veces que aparace la palabra en el texto (palabra fija o que ustedes le escriban la palabra).',
    '•06• Guardar el texto extraído en un archivo de texto.',
    '•07• Cargar el texto del archivo.',
    '•08• Cargar palabras funcionales en español de NLTK.',
    '•09• Tokenizar el texto y eliminar palabras funcionales.',
    '•10• Imprimir algunos detalles sobre los tokens.',
    '•11• Crear un objeto Text de NLTK y calcular la distribución de frecuencia.',
    '•12• Graficar las 40 palabras más comunes.\n',
    '\n○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○ M E N U ○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○\n',
]

print(*Programa, sep="\n")

print('\n► Ingresar un Archivo *.DOCX almacenado en la computadora ("Validar su Ubicación y su formato").\n')

from docx import Document
import os

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def validar_ruta_archivo():
    while True:
        ruta_archivo = input("      Ingresa la ruta del archivo .docx: ")
        # Reemplazar cada barra invertida con dos barras invertidas
        ruta_archivo = ruta_archivo.replace('\\', '\\\\')
        if os.path.exists(ruta_archivo) and ruta_archivo.endswith('.docx'):
            print('\n• Ruta ingresada correctamente •')
            return ruta_archivo
        else:
            print('■ La ruta que ingresaste es incorrecta o no cumple los requerimientos ■\n• Favor de verificar •')

# Validar la ruta del archivo una vez
archivo_docx = validar_ruta_archivo()

def validar_entrada_DocX(file_path):
    try:
        doc = Document(file_path)
        print('• Archivo .docx válido •')
        os.system('pause')
        print('\n• Se mostrará el contenido del Archivo •')
        os.system('\n\npause')
        return doc
    except Exception as e:
        print('■ Error al abrir el archivo ■\n', e)
        return None

# Validar la entrada del archivo .docx
documento = validar_entrada_DocX(archivo_docx)
if documento:
    contenido = read_docx(archivo_docx)
    print("Contenido del documento:")
    print(contenido)
    print('\nEste fue el contenido del archivo cargado.')
    os.system('pause')
    
    # Ruta del archivo TXT donde se guardará el contenido
    archivo_txt = archivo_docx.replace('.docx', '.txt')

    # Guarda el contenido en el archivo TXT
    with open(archivo_txt, 'w') as file:
        file.write(contenido)
        print("-------------------------------------------------------------------------------------------\n")
        print('•01• Enviar el texto del archivo *.DOCX a un archivo de texto.txt almacenado en la computadora.')
        print(f'■ El contenido del archivo DOCX se ha guardado en {archivo_txt} ■')
else:
    print("No se pudo guardar el documento‼")

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")

print('•2• Contar  el numero de palabras:')
# Abrir el archivo de texto
# with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web.txt", "r", encoding="utf-8") as archivo:
with open(archivo_txt) as archivo:
    # Leer el contenido del archivo
    contenido = archivo.read()
    # Dividir el contenido en palabras utilizando el espacio como separador
    palabras = contenido.split()
    # Contar el número de palabras
    num_palabras = len(palabras)

# Imprimir el número de palabras en la página
print("Número de palabras en la página ►", num_palabras)
os.system('pause')
print("-------------------------------------------------------------------------------------------\n")

print('•3• Contar  el numero de lineas:')
# Abrir el archivo de texto
with open(archivo_txt) as archivo:
    # Leer todas las líneas del archivo
    lineas = archivo.readlines()
    # Contar el número de líneas
    num_lineas = len(lineas)

# Imprimir el número de líneas en la página
print("Número de líneas de texto del archivo ►", num_lineas)
os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•4• Mostrar palabras de 3 o 4 caracteres:')
import re

# Abrir el archivo de texto
with open(archivo_txt) as archivo:
    # Leer el contenido del archivo
    contenido = archivo.read()
    # Utilizar una expresión regular para encontrar palabras de 3 o 4 caracteres
    palabras_3_4 = re.findall(r'\b\w{3,4}\b', contenido)

# Imprimir las palabras de 3 o 4 caracteres en 7 columnas
print("Palabras de 3 o 4 caracteres encontradas:")
# Calcular el número de palabras por columna
num_palabras_columna = len(palabras_3_4) // 7
# Imprimir las palabras en 7 columnas
for i in range(num_palabras_columna):
    print("{:15s} {:15s} {:15s} {:15s} {:15s} {:15s} {:15s}".format(
        palabras_3_4[i], palabras_3_4[i + num_palabras_columna], palabras_3_4[i + 2*num_palabras_columna],
        palabras_3_4[i + 3*num_palabras_columna], palabras_3_4[i + 4*num_palabras_columna],
        palabras_3_4[i + 5*num_palabras_columna], palabras_3_4[i + 6*num_palabras_columna]))
# Si hay palabras restantes, imprimirlas en la última columna
if len(palabras_3_4) % 7 != 0:
    for i in range(num_palabras_columna * 7, len(palabras_3_4)):
        print("{:15s}".format(palabras_3_4[i]), end=" ")

print('')
os.system('pause')
print("\n-------------------------------------------------------------------------------------------\n")

print('•5• Cuenta el numero de veces que aparece la palabra en el texto (palabra especificada):')
# Definir la palabra fija que quieres buscar
print('\n■ ■ ■ Ingresa una palabra que quieras utilizar de la lista de palabras de 3 o 4 caracteres encontradas: ■ ■ ■')
palabra_fija = input()

# Abrir el archivo de texto
with open(archivo_txt) as archivo:
    # Leer el contenido del archivo
    contenido = archivo.read()
    # Convertir todo el texto a minúsculas para realizar la búsqueda sin distinción entre mayúsculas y minúsculas
    contenido_minusculas = contenido.lower()
    # Contar el número de veces que aparece la palabra fija en el texto
    apariciones_palabra_fija = contenido_minusculas.count(palabra_fija.lower())

# Imprimir el número de veces que aparece la palabra fija en el texto
print("Número de veces que aparece la palabra •{}• ← {} veces se repite en el documento".format(palabra_fija, apariciones_palabra_fija))
os.system('pause')
print("-------------------------------------------------------------------------------------------\n")

print('•6• Guardar el texto extraído en un archivo de texto:')
# Abrir el archivo de texto original para lectura
with open(archivo_txt) as archivo_origen:
    # Leer el contenido del archivo
    contenido = archivo_origen.read()

# Guardar el texto extraído en un nuevo archivo de texto
nombre_archivo_destino = "texto_extraido.txt"
with open(nombre_archivo_destino, "w") as archivo_destino:
    archivo_destino.write(contenido)
    print(f"El texto extraído se ha guardado correctamente en el archivo ► {nombre_archivo_destino} ◄")

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")

print('•7• Cargar el texto del archivo:')
# Inicializar una variable para almacenar el texto
#texto_extraido = ""

# Abrir el archivo de texto extraído para lectura
with open(nombre_archivo_destino) as archivo_extraido:
    # Leer el contenido del archivo y almacenarlo en la variable
    texto_extraido = archivo_extraido.read()

# Imprimir el texto extraído
print("Texto extraído del archivo ▼")
print(texto_extraido)

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")

print('•8• Cargar palabras funcionales en español de NLTK:')
import nltk

# Descargar palabras funcionales en español de NLTK (si aún no están descargadas)
#nltk.download('stopwords')
print('')
# Cargar palabras funcionales en español
palabras_funcionales = nltk.corpus.stopwords.words("spanish")

# Imprimir algunas palabras funcionales en español
print("Algunas palabras funcionales en español ►")
print(palabras_funcionales[:10])  # Imprime las primeras 10 palabras funcionales
print('')
# Imprimir la cantidad total de palabras funcionales en español
print("Cantidad total de palabras funcionales en español ►", len(palabras_funcionales))

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•9• Tokenizar el texto y eliminar palabras funcionales:')
import nltk

# # Descargar recursos necesarios de NLTK
# nltk.download('punkt')
# nltk.download('stopwords')

# Cargar el texto del archivo
texto_extraido = ""
with open(nombre_archivo_destino) as archivo_extraido:
    texto_extraido = archivo_extraido.read()

# Cargar palabras funcionales en español
palabras_funcionales = nltk.corpus.stopwords.words("spanish")

# Tokenizar el texto
tokens = nltk.word_tokenize(texto_extraido, language="spanish")

# Eliminar palabras funcionales
tokens_limpios = [token for token in tokens if token.lower() not in palabras_funcionales]

# Imprimir algunos detalles sobre los tokens limpios
print("\nTokens limpios  ▼")
num_tokens_por_linea = 10  # Número de tokens por línea
num_lineas = len(tokens_limpios) // num_tokens_por_linea  # Número total de líneas
for i in range(num_lineas + 1):
    inicio = i * num_tokens_por_linea
    fin = (i + 1) * num_tokens_por_linea
    print(tokens_limpios[inicio:fin])  # Imprime un grupo de tokens limpios por línea
os.system('pause')

print("\nNúmero total de tokens ►", len(tokens))
os.system('pause')

print("\nNúmero de tokens limpios ►", len(tokens_limpios))

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")

os.system('pause')
print('•10• Imprimir algunos detalles sobre los tokens:')
# Imprimir algunos detalles sobre los tokens limpios
print("\nAlgunos detalles sobre los tokens limpios ▼")
for i, token in enumerate(tokens_limpios[:20], 1):
    print(f"Token {i}: {token}")

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•11• Crear un objeto Text de NLTK y calcular la distribución de frecuencia.')
os.system('pause')

import nltk
from nltk.probability import FreqDist

# Crear un objeto Text de NLTK
texto_limpio_nltk = nltk.Text(tokens_limpios)

# Calcular la distribución de frecuencia
distribucion_limpia = FreqDist(texto_limpio_nltk)

# Imprimir algunos detalles sobre la distribución de frecuencia
print("\nDistribución de frecuencia de los tokens limpios:")
os.system('pause')
num_palabras_por_linea = 5  # Número de palabras por línea
palabras_comunes = distribucion_limpia.most_common(20)  # Obtener las 20 palabras más comunes
num_lineas = len(palabras_comunes) // num_palabras_por_linea  # Número total de líneas
for i in range(num_lineas + 1):
    inicio = i * num_palabras_por_linea
    fin = (i + 1) * num_palabras_por_linea
    print(palabras_comunes[inicio:fin])  # Imprimir un grupo de palabras por línea

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•12• Graficar las 40 palabras más comunes.\n')
os.system('pause')
import matplotlib.pyplot as plt

# Obtener las palabras y frecuencias
palabras_comunes = distribucion_limpia.most_common(40)

# Filtrar palabras con frecuencias mayores a un umbral
umbral_frecuencia = 3  # Puedes ajustar este umbral según sea necesario
palabras_filtradas = [(palabra, frecuencia) for palabra, frecuencia in palabras_comunes if frecuencia >= umbral_frecuencia]

# Extraer palabras y frecuencias filtradas
palabras_filtradas, frecuencias_filtradas = zip(*palabras_filtradas)

# Graficar las palabras filtradas
plt.figure(figsize=(10, 5))
plt.bar(palabras_filtradas, frecuencias_filtradas)
plt.title("Las palabras más comunes (frecuencia >= {})".format(umbral_frecuencia))
plt.xlabel("Palabra")
plt.ylabel("Frecuencia")
plt.xticks(rotation=45)
plt.grid(True)
plt.show()

print('■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■ FIN del Programa •Documento de WORD - VALIDACION• ■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■\n\n')
os.system('pause')










os.system('cls')
print('\n■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■ INICIO del Programa •Página WEB• ■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■')

Programa=[
'\n○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○\n',
'■ Realizar las siguientes acciones sobre una página web ■',
'► Ingresa una dirección web como: "www.colimanoticias.com"',
'\n'
'•01• Enviar el texto de la pagina a un archivo de "PLN_texto_pagina_web_INPUT_Val.txt" almacenado en la computadora.',
'•02• Contar  el numero de palabras.',
'•03• Contar  el numero de lineas.',
'•04• Mostrar palabras de 3 o 4 caracteres.',
'•05• Cuenta el numero de veces que aparace la palabra en el texto (palabra fija o que ustedes le escriban la palabra).',
'•06• Guardar el texto extraído en un archivo de texto.',
'•07• Cargar el texto del archivo.',
'•08• Cargar palabras funcionales en español de NLTK.',
'•09• Tokenizar el texto y eliminar palabras funcionales.',
'•10• Imprimir algunos detalles sobre los tokens.',
'•11• Crear un objeto Text de NLTK y calcular la distribución de frecuencia.',
'•12• Graficar las 40 palabras más comunes.\n',
'○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○○',
]

print(*Programa, sep = "\n")

#https://www.colimanoticias.com/

import requests

print('•1• Enviar el texto de la pagina a un archivo de texto.txt almacenado en la computadora:')
# URL de la página web
print('Ingresa la Dirección URL completa: ')
url = input()

# Verificar si la URL tiene un esquema válido
if 'http://' not in url and 'https://' not in url:
    url = 'https://' + url

# Obtener el contenido HTML de la página web
response = requests.get(url)

# Verificar si la solicitud fue exitosa
if response.status_code == 200:
    html = response.text
    # Función para limpiar el texto de etiquetas HTML
    def limpiar_html(texto):
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(texto, 'html.parser')
        return soup.get_text()

    # Limpiar el HTML
    texto_pagina = limpiar_html(html)

    # Guardar el texto extraído en un archivo de texto
    if texto_pagina:
        with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val.txt", "w", encoding="utf-8") as archivo_web:
            archivo_web.write(texto_pagina)
            print("El texto de la página se ha guardado correctamente en el archivo ► PLN_texto_pagina_web_INPUT_Val.txt ◄")
    else:
        print("No se pudo extraer texto de la página‼")

else:
    print("No se pudo acceder a la página web. Por favor, verifica la URL e intenta nuevamente.")

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•2• Contar  el numero de palabras:')
os.system('pause')
# Abrir el archivo de texto
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val.txt", "r", encoding="utf-8") as archivo_web:
    # Leer el contenido del archivo
    contenido_web = archivo_web.read()
    # Dividir el contenido en palabras utilizando el espacio como separador
    palabras_web = contenido_web.split()
    # Contar el número de palabras
    num_palabras_web = len(palabras_web)

# Imprimir el número de palabras en la página
print("Número de palabras en la página ►", num_palabras_web)

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•3• Contar  el numero de lineas:')
os.system('pause')
# Abrir el archivo de texto
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val.txt", "r", encoding="utf-8") as archivo_web:
    # Leer todas las líneas del archivo
    lineas_web = archivo_web.readlines()
    # Contar el número de líneas
    num_lineas_web = len(lineas_web)

# Imprimir el número de líneas en la página
print("Número de líneas de texto en la página ►", num_lineas_web)

os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•4• Mostrar palabras de 3 o 4 caracteres:')
import re
os.system('pause')
# Abrir el archivo de texto
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val.txt", "r", encoding="utf-8") as archivo_web:
    # Leer el contenido del archivo
    contenido_web = archivo_web.read()
    # Utilizar una expresión regular para encontrar palabras de 3 o 4 caracteres
    palabras_3_4_web = re.findall(r'\b\w{3,4}\b', contenido_web)

# Imprimir las palabras de 3 o 4 caracteres en 7 columnas
print("Palabras de 3 o 4 caracteres encontradas:")
os.system('pause')
# Calcular el número de palabras por columna
num_palabras_columna_web = len(palabras_3_4_web) // 7
# Imprimir las palabras en 7 columnas
for i in range(num_palabras_columna_web):
    print("{:15s} {:15s} {:15s} {:15s} {:15s} {:15s} {:15s}".format(
        palabras_3_4_web[i], palabras_3_4_web[i + num_palabras_columna_web], palabras_3_4_web[i + 2*num_palabras_columna_web],
        palabras_3_4_web[i + 3*num_palabras_columna_web], palabras_3_4_web[i + 4*num_palabras_columna_web],
        palabras_3_4_web[i + 5*num_palabras_columna_web], palabras_3_4_web[i + 6*num_palabras_columna_web]))
# Si hay palabras restantes, imprimirlas en la última columna
if len(palabras_3_4_web) % 7 != 0:
    for i in range(num_palabras_columna_web * 7, len(palabras_3_4_web)):
        print("{:15s}".format(palabras_3_4_web[i]), end=" ")

os.system('pause')
print("\n-------------------------------------------------------------------------------------------\n")

print('•5• Cuenta el numero de veces que aparece la palabra en el texto (palabra especificada):')
# Definir la palabra fija que quieres buscar
print('\n■ ■ ■ Ingresa una palabra que quieras utilizar de la página web ■ ■ ■')
palabra_fija_web = input()

# Abrir el archivo de texto
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val.txt", "r", encoding="utf-8") as archivo_web:
    # Leer el contenido del archivo
    contenido_web = archivo_web.read()
    # Convertir todo el texto a minúsculas para realizar la búsqueda sin distinción entre mayúsculas y minúsculas
    contenido_minusculas_web = contenido_web.lower()
    # Contar el número de veces que aparece la palabra fija en el texto
    apariciones_palabra_fija_web = contenido_minusculas_web.count(palabra_fija_web.lower())

# Imprimir el número de veces que aparece la palabra fija en el texto
print("Número de veces que aparece la palabra •{}• ← {} veces se repite en el documento".format(palabra_fija_web, apariciones_palabra_fija_web))
os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•6• Guardar el texto extraído en un archivo de texto:')
os.system('pause')
# Abrir el archivo de texto original para lectura
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val.txt", "r", encoding="utf-8") as archivo_origen_web:
    # Leer el contenido del archivo
    contenido_web = archivo_origen_web.read()

# Guardar el texto extraído en un nuevo archivo de texto
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val_EXTRAIDO.txt", "w", encoding="utf-8") as archivo_destino_web:
    archivo_destino_web.write(contenido_web)

print("El texto extraído se ha guardado correctamente en el archivo ► PLN_texto_pagina_web_INPUT_Val_EXTRAIDO.txt ◄")
os.system('pause')
print("-------------------------------------------------------------------------------------------\n")


print('•7• Cargar el texto del archivo:')
os.system('pause')
# Inicializar una variable para almacenar el texto
texto_extraido_web = ""

# Abrir el archivo de texto extraído para lectura
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val_EXTRAIDO.txt", "r", encoding="utf-8") as archivo_extraido_web:
    # Leer el contenido del archivo y almacenarlo en la variable
    texto_extraido_web = archivo_extraido_web.read()

# Imprimir el texto extraído
print("Texto extraído del archivo ▼")
print(texto_extraido_web)
print("-------------------------------------------------------------------------------------------\n")

print('•8• Cargar palabras funcionales en español de NLTK:')
import nltk

# Descargar palabras funcionales en español de NLTK (si aún no están descargadas)
#nltk.download('stopwords')
print('')
# Cargar palabras funcionales en español
palabras_funcionales_web = nltk.corpus.stopwords.words("spanish")

# Imprimir algunas palabras funcionales en español
print("Algunas palabras funcionales en español ►")
print(palabras_funcionales_web[:10])  # Imprime las primeras 10 palabras funcionales
print('')
# Imprimir la cantidad total de palabras funcionales en español
print("Cantidad total de palabras funcionales en español ►", len(palabras_funcionales_web))
print("-------------------------------------------------------------------------------------------\n")


print('•9• Tokenizar el texto y eliminar palabras funcionales:')
import nltk

# # Descargar recursos necesarios de NLTK
# nltk.download('punkt')
# nltk.download('stopwords')

# Cargar el texto del archivo
texto_extraido_web = ""
with open("D:\\Tony\\Documents\\MEGAsync\\6to Sem\\Optativa PLN - Procesamiento de Lenguaje Natural\\PLN\\Actividad_PLN_GitDash\\Documentos\\PLN_texto_pagina_web_INPUT_Val_EXTRAIDO.txt", "r", encoding="utf-8") as archivo_extraido_web:
    texto_extraido_web = archivo_extraido_web.read()

# Cargar palabras funcionales en español
palabras_funcionales_web = nltk.corpus.stopwords.words("spanish")

# Tokenizar el texto
tokens_web = nltk.word_tokenize(texto_extraido_web, language="spanish")

# Eliminar palabras funcionales
tokens_limpios_web = [token for token in tokens_web if token.lower() not in palabras_funcionales_web]

# Imprimir algunos detalles sobre los tokens limpios
print("\nTokens limpios  ▼")
num_tokens_por_linea_web = 10  # Número de tokens por línea
num_lineas_web = len(tokens_limpios_web) // num_tokens_por_linea_web  # Número total de líneas
for i in range(num_lineas_web + 1):
    inicio = i * num_tokens_por_linea_web
    fin = (i + 1) * num_tokens_por_linea_web
    print(tokens_limpios_web[inicio:fin])  # Imprime un grupo de tokens limpios por línea
print("\nNúmero total de tokens ►", len(tokens_web))
print("\nNúmero de tokens limpios ►", len(tokens_limpios_web))
print("-------------------------------------------------------------------------------------------\n")


print('•10• Imprimir algunos detalles sobre los tokens:')
# Imprimir algunos detalles sobre los tokens limpios
print("\nAlgunos detalles sobre los tokens limpios ▼")
for i, token in enumerate(tokens_limpios_web[:20], 1):
    print(f"Token {i}: {token}")
print("-------------------------------------------------------------------------------------------\n")


print('•11• Crear un objeto Text de NLTK y calcular la distribución de frecuencia.')
import nltk
from nltk.probability import FreqDist

# Crear un objeto Text de NLTK
texto_limpio_nltk_web = nltk.Text(tokens_limpios_web)

# Calcular la distribución de frecuencia
distribucion_limpia_web = FreqDist(texto_limpio_nltk_web)

# Imprimir algunos detalles sobre la distribución de frecuencia
print("\nDistribución de frecuencia de los tokens limpios:")
num_palabras_por_linea_web = 5  # Número de palabras por línea
palabras_comunes_web = distribucion_limpia_web.most_common(20)  # Obtener las 20 palabras más comunes
num_lineas_web = len(palabras_comunes_web) // num_palabras_por_linea_web # Número total de líneas
for i in range(num_lineas_web + 1):
    inicio = i * num_palabras_por_linea_web
    fin = (i + 1) * num_palabras_por_linea_web
    print(palabras_comunes_web[inicio:fin])  # Imprimir un grupo de palabras por línea
print("-------------------------------------------------------------------------------------------\n")




print('•12• Graficar las 40 palabras más comunes.\n')
import matplotlib.pyplot as plt

# Obtener las palabras y frecuencias
palabras_comunes_web = distribucion_limpia_web.most_common(40)

# Filtrar palabras con frecuencias mayores a un umbral
umbral_frecuencia_web = 3  # Puedes ajustar este umbral según sea necesario
palabras_filtradas_web = [(palabra, frecuencia) for palabra, frecuencia in palabras_comunes_web if frecuencia >= umbral_frecuencia_web]

# Extraer palabras y frecuencias filtradas
palabras_filtradas_web, frecuencias_filtradas_web = zip(*palabras_filtradas_web)

# Graficar las palabras filtradas
plt.figure(figsize=(10, 5))
plt.bar(palabras_filtradas_web, frecuencias_filtradas_web)
plt.title("Las palabras más comunes (frecuencia >= {})".format(umbral_frecuencia_web))
plt.xlabel("Palabra")
plt.ylabel("Frecuencia")
plt.xticks(rotation=45)
plt.grid(True)
plt.show()

print('■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■ FIN del Programa ■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■■\n\n')
os.system('pause')